"""DOCX 解析器 — 使用 python-docx 提取文本"""

from docx import Document


class DocxParser:
    extension = ".docx"

    def parse(self, file_path: str) -> str:
        """提取 DOCX 全部文本，保留段落和表格"""
        doc = Document(file_path)
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n\n".join(parts) if parts else "（文档内容为空）"
